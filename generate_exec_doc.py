"""Generate the Executive Briefing Document (.docx) for the Model Control team."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page Setup ───────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ───────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x0A, 0x1A, 0x3A)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(8)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(6)
    else:
        hs.font.size = Pt(11.5)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(4)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "0A1A3A")
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")
    return table


def add_diagram(text):
    """Add a monospaced diagram block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONTROL TOWER")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x1A, 0x3A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Unified Model, Tool & GenAI Governance Platform")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x5A)

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Architecture & Governance Overview")
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x7A)

for _ in range(4):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Repository: github.com/fansari100/model-genai-control-tower")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x0A, 0x6A, 0xBA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("February 2026  •  167 Source Files  •  12/12 CI Jobs Green")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x7A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by Farooq Ansari")
run.font.size = Pt(11)
run.font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. System Architecture",
    "3. What the System Governs (Domain Model)",
    "4. Certification Pipeline (Effective Challenge)",
    "5. Risk Rating Engine",
    "6. Compliance Framework Mapping",
    "7. Evaluation Harness (3-Layer Testing)",
    "8. Security & Controls",
    "9. Operational Maturity",
    "10. Technology Stack",
    "11. Repository Structure",
    "12. CI/CD Pipeline Status",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "Control Tower is a production-grade governance platform that provides a single control plane for:"
)
for item in [
    "Maintaining a complete inventory of WM models, bank-impacting models, and non-model tools/EUCs",
    "Producing audit-grade certification evidence for vendor models and GenAI use cases",
    "Running continuous regression, red-teaming, and monitoring with prompt/output logging, model versioning, and agent-risk controls",
]:
    p = doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph(
    "The system operationalizes \"effective challenge\" as defined by SR 11-7 — not as a manual process, "
    "but as a repeatable, policy-enforced, cryptographically auditable pipeline."
)

doc.add_heading("Key Metrics", level=3)
add_table(
    ["Metric", "Value"],
    [
        ["Source Files", "167"],
        ["Lines of Code", "~16,350 (Python 11K, TypeScript 2K, Rego 437, Terraform 692, YAML 2K)"],
        ["Database Tables", "14 (with 15 performance indexes, RLS on 7 tables)"],
        ["API Endpoints", "40+ across 11 route modules"],
        ["CI Pipeline Jobs", "12/12 green (5 CI + 7 SAST)"],
        ["Compliance Frameworks", "7 (SR 11-7, NIST 600-1, OWASP LLM/Agentic, ISO 42001, ATLAS, FINRA)"],
        ["OPA Policies", "4 (approval gates, data classification, agent controls, tool permissions)"],
        ["Evaluation Test Cases", "12+ quality + 15 red-team plugins + 10 PyRIT attacks + 14 garak probes"],
    ],
    col_widths=[2.0, 4.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. System Architecture", level=1)

doc.add_paragraph(
    "The system is organized into five independent planes, each with a distinct responsibility. "
    "This separation ensures that governance decisions, testing, logging, evidence, and policy "
    "enforcement can be audited and scaled independently."
)

doc.add_heading("Five-Plane Architecture", level=2)

add_diagram(
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│                       CONTROL TOWER                            │\n"
    "├──────────────┬──────────────┬──────────┬──────────┬────────────┤\n"
    "│ CONTROL      │ EVALUATION   │ RUNTIME  │ EVIDENCE │ POLICY     │\n"
    "│ PLANE        │ PLANE        │ PLANE    │ PLANE    │ PLANE      │\n"
    "│              │              │          │          │            │\n"
    "│ Inventory +  │ Certification│ Logging, │ Immutable│ Rules,     │\n"
    "│ Governance   │ Tests +      │ Tracing, │ Artifacts│ Gates,     │\n"
    "│ Workflow     │ Red-team     │ Metrics  │ Audit    │ Thresholds │\n"
    "│              │              │          │          │            │\n"
    "│ FastAPI      │ promptfoo    │ OTel     │ MinIO/S3 │ OPA/Rego   │\n"
    "│ Temporal     │ PyRIT        │ Phoenix  │ SHA-256  │ 4 policies │\n"
    "│ PostgreSQL   │ garak        │ Prometheus│Hash chain│            │\n"
    "└──────────────┴──────────────┴──────────┴──────────┴────────────┘"
)

doc.add_heading("Request Flow", level=2)

add_diagram(
    "  [Next.js UI] ──► [API Gateway / Ingress] ──► [FastAPI Backend]\n"
    "                                                  │\n"
    "                    ┌──────────────────────────────┼───────────────┐\n"
    "                    │                              │               │\n"
    "               ┌────▼────┐  ┌──────────┐  ┌───────▼──────┐  ┌────▼────┐\n"
    "               │Inventory│  │Governance │  │ Evaluation   │  │Evidence │\n"
    "               │Service  │  │Workflow   │  │ Orchestrator │  │Service  │\n"
    "               │(CRUD)   │  │(Temporal) │  │ (Workers)    │  │(S3+Hash)│\n"
    "               └────┬────┘  └─────┬─────┘  └──────┬───────┘  └────┬────┘\n"
    "                    │             │               │               │\n"
    "               ┌────▼─────────────▼───────────────▼───────────────▼────┐\n"
    "               │              PostgreSQL 17 + pgvector                 │\n"
    "               │              (14 tables, RLS, 15 indexes)            │\n"
    "               └──────────────────────────────────────────────────────┘\n"
    "                    │             │               │               │\n"
    "               ┌────▼────┐  ┌────▼────┐  ┌───────▼──────┐  ┌────▼────┐\n"
    "               │  OPA    │  │  Kafka  │  │   Phoenix    │  │Prometheus│\n"
    "               │(Policy) │  │ (Audit) │  │ (LLM Trace) │  │(Metrics) │\n"
    "               └─────────┘  └─────────┘  └──────────────┘  └─────────┘"
)

doc.add_heading("Deployment Architecture", level=2)

add_diagram(
    "  Terraform provisions:              Helm deploys onto EKS:\n"
    "  ├─ VPC + 3-AZ subnets             ├─ Backend (3-20 pods, HPA)\n"
    "  ├─ EKS cluster + IRSA             ├─ Frontend (2-10 pods)\n"
    "  ├─ Aurora PostgreSQL 17            ├─ Temporal Worker (2-10 pods)\n"
    "  ├─ S3 + Object Lock (WORM)        ├─ Istio mTLS sidecar\n"
    "  ├─ MSK Kafka (3-6 brokers)        ├─ NetworkPolicy\n"
    "  ├─ KMS encryption keys            ├─ Ingress + TLS\n"
    "  ├─ WAF v2 web ACL                 └─ ServiceMonitor\n"
    "  ├─ ECR registries\n"
    "  └─ IAM roles (IRSA)"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. DOMAIN MODEL
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. What the System Governs (Domain Model)", level=1)

doc.add_paragraph(
    "The domain model covers the three governance mandates of the WM Model Control team: "
    "model inventory, GenAI use case governance, and tool/EUC inventory."
)

doc.add_heading("Entity Relationship Overview", level=2)

add_diagram(
    "  ┌──────────┐     ┌──────────┐     ┌───────────────┐\n"
    "  │  VENDOR  │────▶│  MODEL   │◀────│ GENAI USE CASE│\n"
    "  │          │     │          │     │               │\n"
    "  │ security │     │ type     │     │ risk_rating   │\n"
    "  │ posture  │     │ status   │     │ data_class    │\n"
    "  │ SLA      │     │ risk_tier│     │ uses_agents   │\n"
    "  └──────────┘     └──────────┘     │ uses_rag      │\n"
    "                        │           │ client_facing  │\n"
    "                        │           └───────┬───────┘\n"
    "                        │                   │\n"
    "  ┌──────────┐          │           ┌───────▼───────┐\n"
    "  │   TOOL   │◀─────────┼───────────│  EVAL RUN     │\n"
    "  │          │          │           │  pass_rate    │\n"
    "  │ category │          │           │  OWASP results│\n"
    "  │ attested │          │           └───────┬───────┘\n"
    "  │ criticality│        │                   │\n"
    "  └──────────┘          │           ┌───────▼───────┐\n"
    "                        │           │   FINDING     │\n"
    "                        │           │   severity    │\n"
    "                        │           │   OWASP risk  │\n"
    "                        │           └───────┬───────┘\n"
    "                        │                   │\n"
    "                 ┌──────▼───────┐   ┌───────▼───────┐\n"
    "                 │  APPROVAL    │   │   EVIDENCE    │\n"
    "                 │  decision    │   │   SHA-256     │\n"
    "                 │  hash (tamper│   │   hash chain  │\n"
    "                 │  evident)    │   │   WORM locked │\n"
    "                 └──────────────┘   └───────────────┘"
)

doc.add_heading("Inventory Entities", level=2)

add_table(
    ["Entity", "Purpose", "Key Fields", "MS Equivalent"],
    [
        ["Vendor", "Third-party model/service providers", "security_posture, SLA, red-team due diligence", "Vendor Management"],
        ["Model", "Statistical, ML, LLM, multimodal models", "type, deployment, 7-state status FSM, risk_tier", "Model Inventory"],
        ["Tool", "Non-model EUCs (Excel, VBA, scripts)", "category, criticality, attestation dates", "ClusterSeven IMS"],
        ["GenAIUseCase", "Governed GenAI applications", "risk_rating, architecture flags, OWASP mappings", "GenAI Use Case Registry"],
    ],
    col_widths=[1.2, 1.8, 2.0, 1.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. CERTIFICATION PIPELINE
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. Certification Pipeline (Effective Challenge)", level=1)

doc.add_paragraph(
    "The certification pipeline implements SR 11-7's \"effective challenge\" as a "
    "5-stage, policy-enforced workflow orchestrated by Temporal with human-in-the-loop "
    "approval via signals."
)

add_diagram(
    "  ┌──────────┐   ┌──────────┐   ┌──────────┐   ┌──────────┐   ┌──────────────┐\n"
    "  │ STAGE 0  │──▶│ STAGE 1  │──▶│ STAGE 2  │──▶│ STAGE 3  │──▶│   STAGE 4    │\n"
    "  │          │   │          │   │          │   │          │   │              │\n"
    "  │ Intake & │   │ Doc Pack │   │ Pre-     │   │ Approval │   │ Monitoring & │\n"
    "  │ Risk     │   │ Generate │   │ Deploy   │   │ Gate     │   │ Recertify   │\n"
    "  │ Rating   │   │          │   │ Testing  │   │ (OPA)    │   │              │\n"
    "  └──────────┘   └──────────┘   └──────────┘   └──────────┘   └──────────────┘\n"
    "       │              │              │              │               │\n"
    "   Risk score     AIBOM,        promptfoo,     Human signal    Canary prompts,\n"
    "   Test suites    Model Card    PyRIT, garak   via Temporal    Drift detection,\n"
    "   Committee path                               OPA decision    Recert triggers"
)

doc.add_heading("Certification Pack Output (8 Sections)", level=2)

add_table(
    ["#", "Section", "Content"],
    [
        ["1", "Use Case Summary", "Name, description, risk rating, data classification, architecture flags"],
        ["2", "NIST AI 600-1", "Governance, content provenance, pre-deployment testing, incident disclosure"],
        ["3", "OWASP Mapping", "LLM Top 10 (2025) + Agentic Top 10 (2026) risk-to-control mapping"],
        ["4", "Test Results", "All evaluation runs with pass rates, OWASP category breakdowns"],
        ["5", "Findings Register", "Severity, status, OWASP risk ID, remediation owner/plan/due date"],
        ["6", "Approval Record", "Gate type, decision, approver, rationale, tamper-evident SHA-256 hash"],
        ["7", "Monitoring Plan", "Cadence, canary prompts, thresholds, alert routing, recert triggers"],
        ["8", "ISO 42001 PDCA", "Current lifecycle phase mapping (Plan/Do/Check/Act)"],
    ],
    col_widths=[0.3, 1.5, 4.7],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. RISK RATING ENGINE
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Risk Rating Engine", level=1)

doc.add_paragraph(
    "The risk rating engine uses a weighted factor scoring model. Given use case characteristics, "
    "it deterministically computes a composite risk score that maps to a tier, which determines "
    "required test suites, approval chains, and committee paths."
)

doc.add_heading("Factor Weights", level=2)

add_table(
    ["Factor", "Value", "Weight"],
    [
        ["Data Classification", "RESTRICTED", "50"],
        ["Data Classification", "PII", "40"],
        ["Data Classification", "CONFIDENTIAL", "25"],
        ["Client-Facing", "True", "35"],
        ["Handles PII", "True", "30"],
        ["Uses Agents", "True", "30"],
        ["Category", "AGENT_WORKFLOW", "25"],
        ["Uses Tools", "True", "20"],
        ["Uses Memory", "True", "15"],
        ["Uses RAG", "True", "10"],
    ],
    col_widths=[1.8, 2.0, 0.8],
)

doc.add_heading("Score Thresholds", level=2)

add_table(
    ["Score", "Rating", "Test Suites Required", "Committee Path"],
    [
        ["≥150", "CRITICAL", "9 (all suites incl. PyRIT + garak + agentic)", "WM MRC → Enterprise RC → Board RC"],
        ["≥100", "HIGH", "7 (quality + RAG + security + red-team + controls)", "WM MRC → Enterprise RC"],
        ["≥50", "MEDIUM", "4 (quality + security + red-team + controls)", "WM Model Risk Committee"],
        ["≥20", "LOW", "2 (quality + controls)", "Model Control Review"],
        ["≥0", "MINIMAL", "1 (quality only)", "Model Control Review"],
    ],
    col_widths=[0.6, 1.0, 2.8, 2.1],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. COMPLIANCE MAPPING
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Compliance Framework Mapping", level=1)

doc.add_paragraph(
    "Every regulatory and industry framework requirement is mapped to a specific "
    "control implemented in the system. This is the audit artifact."
)

doc.add_heading("OWASP LLM Top 10 (2025)", level=2)

add_table(
    ["ID", "Risk", "Control Tower Control"],
    [
        ["LLM01", "Prompt Injection", "Cascade guardrails (regex → OpenAI Moderation API) + promptfoo red-team + PyRIT scenarios"],
        ["LLM02", "Insecure Output Handling", "Output guardrails + HITL for client-facing + output logging with PII redaction"],
        ["LLM06", "Sensitive Info Disclosure", "Presidio PII redaction + OPA data_classification policy + retrieval entitlements"],
        ["LLM07", "Excessive Agency", "OPA tool_permissions + per-turn limits + human approval for write operations"],
        ["LLM08", "Data/Model Poisoning", "AIBOM supply-chain transparency + dataset SHA-256 + corpus change → recertification"],
        ["LLM09", "Misinformation", "RAG groundedness evaluation + mandatory citations + canary monitoring"],
    ],
    col_widths=[0.6, 1.3, 4.6],
)

doc.add_heading("OWASP Agentic Top 10 (2026)", level=2)

add_table(
    ["ID", "Risk", "Control Tower Control"],
    [
        ["ASI01", "Agent Goal Hijack", "System prompt hardening + goal drift evaluation suite"],
        ["ASI02", "Tool Misuse", "OPA tool_permissions.rego: strict allowlist + argument schema validation"],
        ["ASI03", "Privilege Abuse", "Scoped credentials per tool + Keycloak IRSA + least-privilege OPA"],
        ["ASI05", "Unexpected RCE", "OPA blocks eval/exec/subprocess patterns + sandboxed execution"],
        ["ASI06", "Memory Poisoning", "Provenance required + TTL enforcement + write count limits (max 10/turn)"],
        ["ASI08", "Cascading Failures", "Temporal retry policies (max 3) + circuit breakers + tool call limits (max 5/turn)"],
        ["ASI10", "Rogue Agents", "Agent registry + signed configs + kill switch (OPA check)"],
    ],
    col_widths=[0.6, 1.3, 4.6],
)

doc.add_heading("Additional Frameworks", level=2)

add_table(
    ["Framework", "Key Mapping"],
    [
        ["SR 11-7 / OCC", "Model Definition → entity model; Effective Challenge → certification pipeline; Governance → committee paths; Monitoring → canaries + drift"],
        ["NIST AI 600-1", "Governance → RBAC + OPA; Content Provenance → citations + AIBOM; Pre-deployment Testing → 3-layer eval; Incident Disclosure → escalation workflow"],
        ["ISO/IEC 42001", "PDCA lifecycle tracked per use case (Plan: intake, Do: test, Check: monitor, Act: remediate)"],
        ["MITRE ATLAS", "Technique IDs on Model and Finding entities (AML.T0020, AML.T0051, etc.)"],
        ["FINRA GenAI", "Prompt/output logging + model version stamping + agentic monitoring"],
    ],
    col_widths=[1.3, 5.2],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. EVALUATION HARNESS
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. Evaluation Harness (3-Layer Testing)", level=1)

add_table(
    ["Layer", "Tool", "Purpose", "Coverage"],
    [
        ["1", "promptfoo", "Quality assertions + adversarial red-teaming", "12 quality tests, 15 attack plugins, 7 delivery strategies"],
        ["2", "PyRIT (Microsoft)", "Systematic GenAI security scenario orchestration", "10 injection vectors with scoring"],
        ["3", "garak (NVIDIA)", "LLM vulnerability scanning", "14 probe categories (inject, leak, hallucinate, encode, toxicity)"],
    ],
    col_widths=[0.5, 1.3, 2.5, 2.2],
)

doc.add_heading("Why Three Layers", level=3)
doc.add_paragraph(
    "promptfoo provides breadth (many test cases, fast execution, CI integration). "
    "PyRIT provides depth (systematic attack trees, multi-turn scenarios). "
    "garak provides coverage (vulnerability categories that the other tools miss). "
    "Three independent evaluation layers is the current institutional standard for GenAI governance."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Security & Controls", level=1)

add_table(
    ["Category", "Implementation"],
    [
        ["Authentication", "Keycloak OIDC, RS256 JWKS verification, 6 RBAC roles"],
        ["Secrets", "HashiCorp Vault (KV v2, Transit encryption, dynamic DB credentials)"],
        ["Guardrails", "2-stage cascade: regex probes → OpenAI Moderation API → human escalation"],
        ["Network", "Istio mTLS (STRICT), K8s NetworkPolicy, AWS WAF v2"],
        ["Data", "PostgreSQL RLS (7 tables), KMS encryption at rest, 15 indexes"],
        ["PII", "Microsoft Presidio (ML-based NER) + regex fallback"],
        ["Audit", "40+ typed events → Kafka (acks=all, idempotent), SHA-256 event hashing"],
        ["Evidence", "Content-addressed storage (SHA-256), hash chains, S3 Object Lock (WORM)"],
        ["SAST/SCA", "Snyk + Trivy + Bandit + Semgrep (7 jobs, all green)"],
    ],
    col_widths=[1.3, 5.2],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. OPERATIONAL MATURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading("9. Operational Maturity", level=1)

doc.add_heading("Service Level Objectives", level=2)

add_table(
    ["SLO", "Target", "Error Budget"],
    [
        ["API Availability", "99.95%", "21.6 min/month"],
        ["API Latency (p99)", "<2 seconds", "Burn-rate alerts at 1h/6h"],
        ["Evaluation Completion", "99%", "Weekly window"],
        ["Evidence Integrity", "100%", "Zero tolerance (P1)"],
        ["Canary Pass Rate", "95%", "Weekly window"],
        ["Audit Event Delivery", "99.99%", "Monthly"],
    ],
    col_widths=[2.0, 1.5, 3.0],
)

doc.add_heading("Disaster Recovery", level=2)

add_table(
    ["Metric", "Value"],
    [
        ["RPO (Recovery Point Objective)", "1 hour"],
        ["RTO (Recovery Time Objective)", "4 hours"],
        ["Primary Region", "us-east-1"],
        ["DR Region", "us-west-2"],
        ["DB Replication", "Aurora Global Database (async, RPO <1s)"],
        ["Evidence Replication", "S3 Cross-Region Replication"],
        ["DR Test Frequency", "Semi-annual"],
    ],
    col_widths=[2.5, 4.0],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
doc.add_heading("10. Technology Stack", level=1)

add_table(
    ["Layer", "Technology", "Version"],
    [
        ["Backend", "Python + FastAPI + SQLAlchemy 2 + Pydantic v2", "3.13"],
        ["Frontend", "Next.js + React + TypeScript + Tailwind CSS", "15 / 19 / 5.7 / 4.0"],
        ["Database", "PostgreSQL + pgvector", "17"],
        ["Workflow", "Temporal", "1.25"],
        ["Policy", "Open Policy Agent + Rego", "0.70"],
        ["Evidence", "MinIO/S3 + SHA-256 hash chains", "WORM Object Lock"],
        ["Events", "Kafka/Redpanda", "KRaft 3.7"],
        ["Auth", "Keycloak (OIDC/JWT)", "26"],
        ["Secrets", "HashiCorp Vault", "KV v2 + Transit"],
        ["Observability", "OpenTelemetry + Phoenix + Prometheus + Grafana", "OTLP"],
        ["Eval", "promptfoo + PyRIT + garak", "Latest"],
        ["Security", "Istio mTLS + WAF + Presidio + Snyk + Trivy", "Latest"],
        ["Deploy", "Terraform + Helm + Docker", "AWS EKS"],
        ["CI/CD", "GitHub Actions", "12 jobs"],
    ],
    col_widths=[1.2, 3.5, 1.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. REPOSITORY STRUCTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading("11. Repository Structure", level=1)

add_diagram(
    "model-genai-control-tower/\n"
    "├── backend/                    Python FastAPI backend\n"
    "│   ├── app/\n"
    "│   │   ├── api/v1/            11 API route modules\n"
    "│   │   ├── models/            14 SQLAlchemy domain models\n"
    "│   │   ├── schemas/           9 Pydantic v2 schema modules\n"
    "│   │   ├── services/          8 service modules (risk, evidence, guardrails...)\n"
    "│   │   ├── security/          Vault + feature flags\n"
    "│   │   ├── integrations/      ServiceNow, Salesforce, LDAP, ClusterSeven\n"
    "│   │   ├── workers/           Eval worker + monitoring worker\n"
    "│   │   ├── utils/             Hashing, PII, PDF, OTEL, logging\n"
    "│   │   ├── auth.py            JWT/OIDC + RBAC\n"
    "│   │   └── main.py            Application factory\n"
    "│   ├── alembic/               2 database migrations\n"
    "│   └── tests/                 Unit + integration + load + E2E\n"
    "├── frontend/                   Next.js 15 frontend\n"
    "│   └── src/app/               11 pages + components + hooks\n"
    "├── workflows/temporal/         Certification workflow + 10 activities\n"
    "├── policies/opa/               4 Rego policy files\n"
    "├── eval/                       promptfoo + PyRIT + garak configs\n"
    "├── deploy/\n"
    "│   ├── terraform/             AWS IaC (VPC, EKS, RDS, S3, MSK, KMS, WAF)\n"
    "│   └── helm/                  Kubernetes deployment chart\n"
    "├── ops/                        SLOs, SOC2, alerting, runbooks\n"
    "├── security/                   Istio mTLS, WAF rules\n"
    "├── scripts/                    Seed data, AIBOM gen, cert pack gen\n"
    "└── .github/workflows/          4 CI/CD pipelines (12 jobs)"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. CI/CD STATUS
# ══════════════════════════════════════════════════════════════
doc.add_heading("12. CI/CD Pipeline Status", level=1)

doc.add_paragraph("All 12 CI/CD jobs pass on the latest commit:")

doc.add_heading("CI Pipeline (5 jobs)", level=2)

add_table(
    ["Job", "Status", "What It Verifies"],
    [
        ["Backend Lint", "✅ PASS", "ruff check + ruff format + mypy type checking"],
        ["Backend Tests", "✅ PASS", "32 tests against PostgreSQL 17 via Alembic migrations"],
        ["Frontend Build", "✅ PASS", "TypeScript compilation + Next.js 15 production build"],
        ["OPA Policy Tests", "✅ PASS", "opa check on all 4 Rego policy files"],
        ["Docker Build", "✅ PASS", "Multi-stage backend + frontend image builds"],
    ],
    col_widths=[1.5, 0.8, 4.2],
)

doc.add_heading("Security SAST Pipeline (7 jobs)", level=2)

add_table(
    ["Job", "Status", "What It Scans"],
    [
        ["Snyk Python", "✅ PASS", "Python dependency vulnerabilities"],
        ["Snyk Node", "✅ PASS", "Node.js dependency vulnerabilities"],
        ["Trivy Backend", "✅ PASS", "Backend container image CVEs → GitHub SARIF"],
        ["Trivy Frontend", "✅ PASS", "Frontend container image CVEs → GitHub SARIF"],
        ["Semgrep SAST", "✅ PASS", "Multi-language static analysis patterns"],
        ["Bandit SAST", "✅ PASS", "Python security issue detection"],
        ["OPA Validation", "✅ PASS", "Policy syntax + logic validation"],
    ],
    col_widths=[1.5, 0.8, 4.2],
)

# ── Footer ───────────────────────────────────────────────────
doc.add_paragraph("")
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x8A, 0x8A, 0xAA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Repository: github.com/fansari100/model-genai-control-tower")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x0A, 0x6A, 0xBA)

# ── Save ─────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "Control_Tower_Executive_Briefing.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
